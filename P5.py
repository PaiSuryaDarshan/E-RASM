# Un-Licensed. Free to use, copy, modify and merge.
# Please maintain attribution to the original author and the license.

from docx import Document

CHEMICAL_NUMBER = 0
CHEMICAL_NAME   = "Ethanol"
MOL_WEIGHT      = ""
DENSITY         = ""
QUANTITY        = ""
MMOL            = ""
EQUIVALENT      = ""
interpreted_haz_list_demo = ['', '', '\n2', '', '', '\nX', '', '', '']
OTHERS_TO_SPECIFY = ""

# First Step: 
# Preparing list that covers the full length of the table (Final Touches)

def prepare_final_list(
    interpreted_haz_list: list, 
    CHEMICAL_NAME = "", 
    MOL_WEIGHT = "", 
    DENSITY = "", 
    QUANTITY = "", 
    MMOL = "", 
    EQUIVALENT = "", 
    BLANK_ENTRY = "",
    OTHERS_TO_SPECIFY = ""
    ) -> list:

    final_entry = []
    final_entry.append(CHEMICAL_NAME)
    final_entry.append(MOL_WEIGHT)

    # remove decimal point, (if present)
    DENSITY_STRIPED = DENSITY.replace('.', '')
    if DENSITY_STRIPED.isnumeric():
        final_entry.append(DENSITY)
    else:
        final_entry.append("")
    
    final_entry.append(QUANTITY)
    final_entry.append(MMOL)
    final_entry.append(EQUIVALENT)
    final_entry.append(BLANK_ENTRY)

    for haz_tag in interpreted_haz_list:
        final_entry.append(haz_tag)
    
    final_entry.append(OTHERS_TO_SPECIFY)

    return final_entry

# Second Step: 
# Entering value into the document

# ----------------------------------------------------------------
# Perquisites for step two
# ----------------------------------------------------------------
# function that deletes rows 
def delete_row_in_table(table, row):
    table._tbl.remove(table.rows[row]._tr)
    return
# ----------------------------------------------------------------


def make_table_entry(
        CHEMICAL_NUMBER, 
        final_entry,
        haz_table 
        ):

    # MATCHES CHEMICAL NUMBER WITH TABLE ROW INDEX
    row_index = CHEMICAL_NUMBER + 3
    for i in range(len(final_entry)):
        haz_table.rows[row_index].cells[i].text = final_entry[i]




# TOTAL NUMBER OF ROWS IN TEMPLATE DOCUMENT, HAZARD TABLE: 25
# Del unfilled rows
def del_excess_rows(NAME = "test_draft"):
    # delete all empty rows
    path = f'./Final Docs/{NAME}.docx'
    document = Document(path)
    haz_table = document.tables[0]
    for i in range(len(haz_table.rows)-1, 0, -1):
        if haz_table.rows[i].cells[0].text.strip() == "":
            delete_row_in_table(haz_table, i)

    # Save in the same path as original, This overwrites the document
    document.save(path)
















# TEST
def reference():
    document = Document('./template_directory/Blank Reaction Risk Assessment Form.docx')

    haz_table = document.tables[0]
    # For developer's reference
    for i in range(0,6):
        print(i, haz_table.rows[2].cells[i].text)
        continue
    for i in range(6,16):
        print(i, haz_table.rows[1].cells[i].text)
        continue

    print(f"TOTAL NUMBER OF ROWS IN HAZARD TABLE OF TEMPLATE DOCUMENT (INCLUDING TWO ROWS FOR HEADER): {len(haz_table.rows)}")
    print(f"TOTAL NUMBER OF ROWS IN HAZARD TABLE OF TEMPLATE DOCUMENT (EXCLUDING TWO ROWS FOR HEADER): {len(haz_table.rows)-2}")

if __name__ == "__main__":
    # reference()
    # final_entry_demo = prepare_final_list(interpreted_haz_list_demo, CHEMICAL_NAME)
    # make_table_entry(CHEMICAL_NUMBER, final_entry_demo)
    # del_excess_rows()
    pass